"""
AI機能クライアントモジュール (CAD→概算見積もり生成AI / 日本住建株式会社)

Streamlitアプリの3つのAI機能の基盤を提供する:
  (A) 音声データの文字起こし (打合せ音声・音声コマンド) - OpenAI Whisper API
  (B) 提案書・打合せ記録からの見積変更案抽出 - Anthropic Messages API
  (C) 音声コマンドテキストの数量変更指示への構造化 - Anthropic Messages API

設定 (環境変数 → Streamlit secrets の順で取得):
  - ANTHROPIC_API_KEY : テキスト解析用 (必須)
  - OPENAI_API_KEY    : Whisper文字起こし用 (任意)
  - AI_MODEL          : 使用モデル (デフォルト 'claude-sonnet-4-6')

依存: 標準ライブラリ + requests + pypdf + python-docx + openpyxl
"""

from __future__ import annotations

import io
import json
import os
from typing import Any

import requests

# --- 定数 ---------------------------------------------------------------

ANTHROPIC_ENDPOINT = "https://api.anthropic.com/v1/messages"
ANTHROPIC_VERSION = "2023-06-01"
OPENAI_TRANSCRIPTION_ENDPOINT = "https://api.openai.com/v1/audio/transcriptions"

DEFAULT_MODEL = "claude-sonnet-4-6"
WHISPER_MODEL = "whisper-1"

TEXT_TIMEOUT_SEC = 120   # テキスト解析 (Anthropic)
AUDIO_TIMEOUT_SEC = 300  # 音声文字起こし (Whisper)
MAX_TOKENS = 8000

# 変更案スキーマの許容値
_CHANGE_TYPES = {"price", "quantity", "spec", "add", "remove"}
_CONFIDENCES = {"high", "medium", "low"}

# 音声ファイル拡張子 → MIMEタイプ (Whisper対応形式)
_AUDIO_MIME_TYPES = {
    ".wav": "audio/wav",
    ".mp3": "audio/mpeg",
    ".m4a": "audio/mp4",
    ".mp4": "audio/mp4",
    ".webm": "audio/webm",
    ".ogg": "audio/ogg",
    ".oga": "audio/ogg",
    ".flac": "audio/flac",
    ".mpga": "audio/mpeg",
    ".mpeg": "audio/mpeg",
}


# --- 例外 ---------------------------------------------------------------

class AIClientError(Exception):
    """AI機能の呼び出しに失敗したときに送出する例外 (メッセージは日本語)。"""


# --- 内部ヘルパー -------------------------------------------------------

def _get_secret(key: str) -> str:
    """環境変数 → Streamlit secrets の順で値を取得する。どちらにも無ければ空文字を返す。"""
    val = os.environ.get(key)
    if val:
        return val
    try:
        import streamlit as st  # type: ignore

        # st.secrets はファイル未配置時に StreamlitSecretNotFoundError を投げる
        return str(st.secrets.get(key, "")) if hasattr(st, "secrets") else ""
    except Exception:  # noqa: BLE001
        return ""


def _get_model() -> str:
    """使用するAnthropicモデルIDを返す。"""
    return _get_secret("AI_MODEL") or DEFAULT_MODEL


def _call_anthropic(system: str, user_text: str) -> str:
    """Anthropic Messages API を requests 直呼びで実行し、応答テキストを返す。"""
    api_key = _get_secret("ANTHROPIC_API_KEY")
    if not api_key:
        raise AIClientError(
            "ANTHROPIC_API_KEY が設定されていません。環境変数または Streamlit secrets に設定してください。"
        )
    try:
        resp = requests.post(
            ANTHROPIC_ENDPOINT,
            json={
                "model": _get_model(),
                "max_tokens": MAX_TOKENS,
                "system": system,
                "messages": [{"role": "user", "content": user_text}],
            },
            headers={
                "x-api-key": api_key,
                "anthropic-version": ANTHROPIC_VERSION,
                "Content-Type": "application/json",
            },
            timeout=TEXT_TIMEOUT_SEC,
        )
    except requests.Timeout as exc:
        raise AIClientError(
            f"AI解析がタイムアウトしました ({TEXT_TIMEOUT_SEC}秒)。時間をおいて再試行してください。"
        ) from exc
    except requests.RequestException as exc:
        raise AIClientError(f"AI解析APIへの接続に失敗しました: {exc}") from exc

    if resp.status_code >= 300:
        raise AIClientError(
            f"AI解析APIがエラーを返しました (HTTP {resp.status_code}): {resp.text[:300]}"
        )

    body = resp.json()
    texts = [
        block.get("text", "")
        for block in body.get("content", [])
        if block.get("type") == "text"
    ]
    text = "".join(texts).strip()
    if not text:
        raise AIClientError("AI解析APIの応答にテキストが含まれていませんでした。")
    return text


def _strip_code_fence(text: str) -> str:
    """応答テキストから Markdown コードフェンス (```json 〜 ```) を除去する。"""
    stripped = text.strip()
    if stripped.startswith("```"):
        # 先頭行 (``` または ```json) を落とす
        lines = stripped.splitlines()
        lines = lines[1:]
        # 末尾の ``` を落とす
        while lines and lines[-1].strip() in ("```", ""):
            lines.pop()
        stripped = "\n".join(lines).strip()
    return stripped


def _parse_changes_json(text: str) -> list[dict[str, Any]]:
    """応答テキストを変更案リストとして json.loads する。失敗時は例外を送出する。"""
    parsed = json.loads(_strip_code_fence(text))
    if isinstance(parsed, dict):
        # {"changes": [...]} のようにラップされた場合は中のリストを取り出す
        for value in parsed.values():
            if isinstance(value, list):
                parsed = value
                break
        else:
            raise ValueError("JSONが配列形式ではありません")
    if not isinstance(parsed, list):
        raise ValueError("JSONが配列形式ではありません")
    return [item for item in parsed if isinstance(item, dict)]


def _normalize_change(item: dict[str, Any]) -> dict[str, Any]:
    """変更案1件を規定スキーマに正規化する。"""
    change_type = str(item.get("change_type", "")).strip().lower()
    if change_type not in _CHANGE_TYPES:
        change_type = "spec"
    confidence = str(item.get("confidence", "")).strip().lower()
    if confidence not in _CONFIDENCES:
        confidence = "low"
    return {
        "category": str(item.get("category", "") or "").strip(),
        "item_name": str(item.get("item_name", "") or "").strip(),
        "change_type": change_type,
        "current": str(item.get("current", "") or "").strip(),
        "proposed": str(item.get("proposed", "") or "").strip(),
        "reason": str(item.get("reason", "") or "").strip(),
        "confidence": confidence,
    }


_SCHEMA_INSTRUCTION = """\
出力は必ずJSON配列のみとし、説明文・前置き・コードフェンスは一切付けないこと。
各要素は以下のキーを持つオブジェクトとする:
  - "category"   : 対象の工事種別名 (例: "住宅設備機器工事")
  - "item_name"  : 対象の明細名 (例: "システムバス")
  - "change_type": 変更の種類。"price" (単価・金額の変更) / "quantity" (数量の変更) /
                   "spec" (仕様・型番・グレードの変更) / "add" (明細の追加) / "remove" (明細の削除) のいずれか
  - "current"    : 現在の見積上の内容 (数量・単価・仕様など。不明なら空文字)
  - "proposed"   : 変更後の内容
  - "reason"     : 根拠となる記述の引用 (原文からそのまま引用する)
  - "confidence" : 確度。"high" / "medium" / "low" のいずれか。
                   見積サマリ内の明細と明確に対応付けられる場合は "high"、
                   対応する明細が推測になる場合は "medium"、
                   変更対象が特定できず候補として提示する場合は "low"
変更が読み取れない場合は空配列 [] を返すこと。"""

_RETRY_INSTRUCTION = (
    "\n\n注意: 前回の応答はJSONとして解析できませんでした。"
    "今回は必ず有効なJSON配列だけを出力してください。"
)


def _run_change_extraction(system: str, user_text: str) -> list[dict[str, Any]]:
    """変更案抽出プロンプトを実行し、JSON解析に失敗した場合は1回だけ再試行する。"""
    raw = _call_anthropic(system, user_text)
    try:
        items = _parse_changes_json(raw)
    except (json.JSONDecodeError, ValueError):
        # 1回だけ再試行 (より厳格な指示を追加)
        raw = _call_anthropic(system, user_text + _RETRY_INSTRUCTION)
        try:
            items = _parse_changes_json(raw)
        except (json.JSONDecodeError, ValueError) as exc:
            raise AIClientError(
                f"AIの応答をJSONとして解析できませんでした: {raw[:200]}"
            ) from exc
    return [_normalize_change(item) for item in items]


# --- 公開 API -----------------------------------------------------------

def ai_capabilities() -> dict:
    """利用可能なAI機能を返す。

    Returns:
        dict: {
            "text": bool,        # ANTHROPIC_API_KEY があればテキスト解析可
            "transcribe": bool,  # OPENAI_API_KEY があれば文字起こし可
        }
    """
    return {
        "text": bool(_get_secret("ANTHROPIC_API_KEY")),
        "transcribe": bool(_get_secret("OPENAI_API_KEY")),
    }


def transcribe_audio(audio_bytes: bytes, filename: str) -> str:
    """音声データを OpenAI Whisper API で日本語文字起こしする。

    Args:
        audio_bytes: 音声ファイルのバイナリデータ。
        filename: 元ファイル名 (拡張子からMIMEタイプを判定する)。

    Returns:
        str: 文字起こし結果のテキスト。

    Raises:
        AIClientError: APIキー未設定・通信エラー・APIエラーの場合。
    """
    api_key = _get_secret("OPENAI_API_KEY")
    if not api_key:
        raise AIClientError(
            "OPENAI_API_KEY が設定されていません。音声文字起こしを利用するには設定してください。"
        )
    if not audio_bytes:
        raise AIClientError("音声データが空です。")

    ext = os.path.splitext(filename)[1].lower()
    mime = _AUDIO_MIME_TYPES.get(ext, "application/octet-stream")

    try:
        resp = requests.post(
            OPENAI_TRANSCRIPTION_ENDPOINT,
            headers={"Authorization": f"Bearer {api_key}"},
            files={"file": (filename, audio_bytes, mime)},
            data={
                "model": WHISPER_MODEL,
                "language": "ja",
                "response_format": "json",
            },
            timeout=AUDIO_TIMEOUT_SEC,
        )
    except requests.Timeout as exc:
        raise AIClientError(
            f"文字起こしがタイムアウトしました ({AUDIO_TIMEOUT_SEC}秒)。音声を短く分割して再試行してください。"
        ) from exc
    except requests.RequestException as exc:
        raise AIClientError(f"文字起こしAPIへの接続に失敗しました: {exc}") from exc

    if resp.status_code >= 300:
        raise AIClientError(
            f"文字起こしAPIがエラーを返しました (HTTP {resp.status_code}): {resp.text[:300]}"
        )

    text = str(resp.json().get("text", "")).strip()
    if not text:
        raise AIClientError("文字起こし結果が空でした。音声の内容をご確認ください。")
    return text


def analyze_document_changes(doc_text: str, estimate_summary: str) -> list[dict]:
    """提案書・打合せ記録テキストから「見積のどこが変わるか」の変更案を抽出する。

    Args:
        doc_text: 提案書・打合せ記録などの本文テキスト。
        estimate_summary: 現在の見積サマリ (工事種別と主要明細のJSON文字列)。

    Returns:
        list[dict]: 変更案のリスト。各要素は
            {"category", "item_name", "change_type", "current",
             "proposed", "reason", "confidence"} のキーを持つ。

    Raises:
        AIClientError: APIキー未設定・通信エラー・JSON解析失敗 (再試行後) の場合。
    """
    if not (doc_text or "").strip():
        raise AIClientError("解析対象のテキストが空です。")

    system = (
        "あなたは住宅建築の見積担当者を支援するAIアシスタントです。"
        "提案書や打合せ記録を読み、現在の見積のどの項目がどう変わるかを正確に抽出します。"
        "文書に書かれていない変更を創作してはいけません。\n\n"
        + _SCHEMA_INSTRUCTION
    )
    user_text = (
        "## 現在の見積サマリ (JSON)\n"
        f"{estimate_summary}\n\n"
        "## 提案書・打合せ記録\n"
        f"{doc_text}\n\n"
        "上記の文書から、見積に反映すべき変更点をすべて抽出し、JSON配列で出力してください。"
    )
    return _run_change_extraction(system, user_text)


def parse_voice_command(transcript: str, estimate_summary: str) -> list[dict]:
    """音声コマンドの文字起こしテキストを数量変更指示に構造化する。

    例: 「システムバスを1.25坪タイプに変更して」「エアコンを2台に減らして」
        「キッチンの見積単価を150万に」

    Args:
        transcript: 音声コマンドの文字起こしテキスト。
        estimate_summary: 現在の見積サマリ (工事種別と主要明細のJSON文字列)。

    Returns:
        list[dict]: analyze_document_changes と同じスキーマの変更案リスト。
            変更対象が特定できない場合は confidence='low' の候補を返す。

    Raises:
        AIClientError: APIキー未設定・通信エラー・JSON解析失敗 (再試行後) の場合。
    """
    if not (transcript or "").strip():
        raise AIClientError("音声コマンドのテキストが空です。")

    system = (
        "あなたは住宅建築の見積担当者を支援するAIアシスタントです。"
        "担当者が話した音声コマンドの文字起こしを読み、見積への変更指示として構造化します。"
        "口語表現・言い間違い・単位の省略 (例: 「150万」=1,500,000円) を適切に解釈してください。"
        "指示に含まれない変更を創作してはいけません。\n\n"
        + _SCHEMA_INSTRUCTION
        + "\n変更対象の明細が見積サマリから特定できない場合も、最も近い候補を "
        '"confidence": "low" として返すこと。'
    )
    user_text = (
        "## 現在の見積サマリ (JSON)\n"
        f"{estimate_summary}\n\n"
        "## 音声コマンド (文字起こし)\n"
        f"{transcript}\n\n"
        "上記の音声コマンドを見積への変更指示として構造化し、JSON配列で出力してください。"
    )
    return _run_change_extraction(system, user_text)


def extract_text_from_upload(file_bytes: bytes, filename: str) -> str:
    """アップロードされたファイルからテキストを抽出する。

    対応形式: .pdf (pypdf) / .docx (python-docx) / .xlsx (openpyxl) / .txt

    Args:
        file_bytes: ファイルのバイナリデータ。
        filename: 元ファイル名 (拡張子で処理を分岐する)。

    Returns:
        str: 抽出したテキスト。

    Raises:
        ValueError: 未対応の拡張子の場合。
        AIClientError: ファイルの読み取りに失敗した場合。
    """
    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            pages = [(page.extract_text() or "") for page in reader.pages]
            return "\n\n".join(p.strip() for p in pages if p.strip())

        if ext == ".docx":
            import docx

            document = docx.Document(io.BytesIO(file_bytes))
            parts: list[str] = [
                para.text for para in document.paragraphs if para.text.strip()
            ]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append("\t".join(cells))
            return "\n".join(parts)

        if ext == ".xlsx":
            import openpyxl

            workbook = openpyxl.load_workbook(
                io.BytesIO(file_bytes), read_only=True, data_only=True
            )
            parts = []
            for sheet in workbook.worksheets:
                parts.append(f"## シート: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    cells = ["" if v is None else str(v) for v in row]
                    if any(c.strip() for c in cells):
                        parts.append("\t".join(cells))
            workbook.close()
            return "\n".join(parts)

        if ext == ".txt":
            for encoding in ("utf-8", "cp932", "shift_jis"):
                try:
                    return file_bytes.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return file_bytes.decode("utf-8", errors="replace")

    except (ValueError, AIClientError):
        raise
    except Exception as exc:  # noqa: BLE001 - ライブラリ例外をまとめて変換
        raise AIClientError(
            f"ファイル「{filename}」の読み取りに失敗しました: {type(exc).__name__}: {exc}"
        ) from exc

    raise ValueError(
        f"未対応のファイル形式です: {ext or '(拡張子なし)'} "
        "(対応形式: .pdf / .docx / .xlsx / .txt)"
    )


# --- CLI エントリポイント -----------------------------------------------

if __name__ == "__main__":
    print("=== ai_client.py 動作確認 ===")
    print(f"利用可能な機能: {ai_capabilities()}")
